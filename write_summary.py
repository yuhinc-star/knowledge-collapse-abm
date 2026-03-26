"""
Generate project summary as a formatted .docx
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

OUT = "/Users/jacob/Desktop/Claude Code/Knowledge_Collapse_Project_Summary.docx"

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin   = Inches(1.25)
    section.right_margin  = Inches(1.25)

# ── Helpers ───────────────────────────────────────────────────────────────────
BLUE   = RGBColor(0x1E, 0x40, 0xAF)   # dark blue headings
GREY   = RGBColor(0x37, 0x41, 0x51)   # dark grey
RED    = RGBColor(0xB9, 0x1C, 0x1C)   # emphasis red
GREEN  = RGBColor(0x06, 0x59, 0x37)   # result green
BLACK  = RGBColor(0x11, 0x18, 0x27)

def set_run(run, bold=False, italic=False, size=11, color=BLACK, font="Calibri"):
    run.bold   = bold
    run.italic = italic
    run.font.size  = Pt(size)
    run.font.color.rgb = color
    run.font.name  = font

def heading(text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18 if level == 1 else 12)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    set_run(run, bold=True,
            size=18 if level == 1 else (14 if level == 2 else 12),
            color=BLUE if level <= 2 else GREY)
    return p

def body(text, space_after=4, italic=False, color=BLACK, bold=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after  = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    run = p.add_run(text)
    set_run(run, bold=bold, italic=italic, color=color)
    return p

def bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after  = Pt(2)
    p.paragraph_format.left_indent  = Inches(0.3)
    if bold_prefix:
        r1 = p.add_run(bold_prefix + ": ")
        set_run(r1, bold=True, size=11)
        r2 = p.add_run(text)
        set_run(r2, size=11)
    else:
        run = p.add_run(text)
        set_run(run, size=11)
    return p

def rule():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(6)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "BFDBFE")
    pBdr.append(bottom)
    pPr.append(pBdr)

def add_table(headers, rows, col_widths=None):
    """Add a styled table with header row."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    # header row
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        set_run(run, bold=True, size=10, color=RGBColor(0xFF,0xFF,0xFF))
        shading = OxmlElement("w:shd")
        shading.set(qn("w:val"),   "clear")
        shading.set(qn("w:color"), "auto")
        shading.set(qn("w:fill"),  "1E40AF")
        cell._tc.get_or_add_tcPr().append(shading)
    # data rows
    for ri, row_data in enumerate(rows):
        row = table.rows[ri + 1]
        fill = "EFF6FF" if ri % 2 == 0 else "FFFFFF"
        for ci, val in enumerate(row_data):
            cell = row.cells[ci]
            cell.text = ""
            run = cell.paragraphs[0].add_run(str(val))
            set_run(run, size=10)
            shading = OxmlElement("w:shd")
            shading.set(qn("w:val"),   "clear")
            shading.set(qn("w:color"), "auto")
            shading.set(qn("w:fill"),  fill)
            cell._tc.get_or_add_tcPr().append(shading)
    # column widths
    if col_widths:
        for ci, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[ci].width = Inches(w)
    doc.add_paragraph()   # spacing after table


# ═════════════════════════════════════════════════════════════════════════════
# TITLE PAGE
# ═════════════════════════════════════════════════════════════════════════════
tp = doc.add_paragraph()
tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
tp.paragraph_format.space_before = Pt(48)
r = tp.add_run("Knowledge Collapse")
set_run(r, bold=True, size=28, color=BLUE)

tp2 = doc.add_paragraph()
tp2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = tp2.add_run("AI, Human Cognition, and the Erosion of Shared Knowledge")
set_run(r2, italic=True, size=14, color=GREY)

tp3 = doc.add_paragraph()
tp3.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = tp3.add_run("Project Summary  ·  Based on Acemoglu, Kong & Ozdaglar (2026)")
set_run(r3, size=11, color=GREY)

doc.add_paragraph()
rule()
doc.add_paragraph()


# ═════════════════════════════════════════════════════════════════════════════
# PART 1 — PLAIN LANGUAGE SUMMARY
# ═════════════════════════════════════════════════════════════════════════════
heading("Part 1 — Plain Language Summary", level=1)

body(
    "This project starts with an academic economics paper and turns it into something "
    "you can interact with: a visual simulation where you move sliders and watch how "
    "societies learn — or fail to learn — when AI becomes widely available.",
    space_after=8
)

heading("The core idea", level=2)
body(
    "Imagine a community of people who all need to know two things to do their jobs well: "
    "(1) a shared body of knowledge that everyone relies on — think domain expertise, "
    "scientific consensus, professional standards — and (2) the specific details of "
    "their own individual task or context.",
    space_after=6
)
body(
    "The only way shared knowledge stays current is if people actively engage with it. "
    "Every time someone does the hard work of learning, they inadvertently contribute "
    "a little to the community's collective understanding — a by-product of their effort. "
    "This is a classic public good: everyone benefits, but no one has a strong private "
    "reason to produce it.",
    space_after=6
)
body(
    "Now introduce an AI assistant. The AI is very good at telling each person exactly "
    "what they need to know about their specific task. So people stop doing the hard work "
    "themselves. Individually, each person is better off. But the by-product — the "
    "contribution to shared knowledge — disappears. Over time, the community's shared "
    "knowledge degrades. And because shared and task-specific knowledge are complements "
    "(you need both to do anything useful), everyone eventually becomes worse off.",
    space_after=6
)
body(
    "This is knowledge collapse: a self-reinforcing spiral where AI substitutes for "
    "human learning, shared knowledge erodes, and the value of both human effort and "
    "AI assistance falls together.",
    space_after=8, bold=True
)

heading("What we built", level=2)
body(
    "We translated the paper's mathematical model into a running simulation and built "
    "an interactive web application with five panels:",
    space_after=4
)
bullet("Knowledge Dynamics — shows the mathematical law governing how shared knowledge evolves period to period, and where the stable and unstable equilibria are")
bullet("Time Paths — shows trajectories of knowledge over time from different starting points, illustrating path dependence and the tipping point")
bullet("Social Welfare — shows how total societal welfare changes as AI capability increases, revealing the hump-shaped curve with an interior optimum")
bullet("Community Size Effect — shows how larger communities are more resilient, and why the protection only scales logarithmically")
bullet("Agent Beliefs — shows individual agents as dots in belief space, visualising how tightly or loosely they track the truth in both shared and task-specific dimensions")
doc.add_paragraph()

heading("The most important findings", level=2)
bullet("There are two fundamentally different regimes", bold_prefix="Two worlds")
body("    If people's willingness to learn is sufficiently elastic (responds strongly "
     "to incentives), collapse is possible. If it is inelastic, shared knowledge "
     "always recovers. A single parameter — the curvature of effort costs — determines which world you are in.",
     space_after=4)

bullet("The danger is non-obvious", bold_prefix="Silent erosion")
body("    For most of the range of AI capability, the tipping point (the minimum knowledge "
     "level from which recovery is possible) is negligibly small. Society appears safe. "
     "Then, near the collapse threshold, the tipping point grows rapidly and the high-knowledge "
     "equilibrium shrinks — a sudden cliff, not a gradual slope.",
     space_after=4)

bullet("More AI is not always better", bold_prefix="Non-monotone welfare")
body("    Welfare is hump-shaped in AI capability. A little AI helps; too much destroys "
     "the conditions that make AI useful in the first place. There is an interior "
     "welfare-maximising level that a regulator should target.",
     space_after=4)

bullet("Size helps, but not enough", bold_prefix="Logarithmic scaling")
body("    Larger communities (universities, professional networks, institutions) have "
     "higher collapse thresholds. But the protection grows only as the logarithm of size — "
     "you need exponentially more people to withstand linearly better AI.",
     space_after=8)

heading("A surprising result from the simulation", level=2)
body(
    "In the Agent Beliefs panel, the gap between what any individual knows and what "
    "the community collectively knows is largest when society is doing well — not when "
    "it is struggling. In the high-knowledge equilibrium, agents work hard, produce "
    "rich collective signals, and the aggregate far exceeds any individual. The public-good "
    "problem is not a symptom of decline; it is a structural feature of success. "
    "This makes it invisible until it is too late.",
    space_after=8, italic=True
)

rule()
doc.add_page_break()


# ═════════════════════════════════════════════════════════════════════════════
# PART 2 — DETAILED TECHNICAL SUMMARY
# ═════════════════════════════════════════════════════════════════════════════
heading("Part 2 — Detailed Technical Summary", level=1)

# ── 2.1 The paper ─────────────────────────────────────────────────────────────
heading("2.1  The Source Paper", level=2)
body(
    "Acemoglu, Kong & Ozdaglar (2026), 'AI, Human Cognition and Knowledge Collapse.' "
    "The paper constructs a formal dynamic model of knowledge production in an economy "
    "where agents simultaneously produce private and public signals through effortful learning. "
    "It characterises conditions under which the introduction of AI assistance triggers "
    "permanent loss of general knowledge.",
    space_after=6
)

# ── 2.2 Model structure ───────────────────────────────────────────────────────
heading("2.2  Mathematical Model Structure", level=2)

heading("States of the world", level=3)
body("The model has two knowledge dimensions per period t:", space_after=2)
bullet("θ_t  —  the common/shared knowledge state, following a random walk: "
       "θ_{t+1} = θ_t + ε_t,  ε_t ~ N(0, Σ²). It drifts and must be continuously refreshed.",
       bold_prefix="Common state")
bullet("θ_{i,t}  —  the idiosyncratic/task-specific state for agent i, drawn fresh "
       "each period from N(0,1). It cannot be shared or accumulated.",
       bold_prefix="Idiosyncratic state")
doc.add_paragraph()

heading("Production", level=3)
body(
    "Output has a Leontief (perfect-complement) structure: "
    "Output = G(X_t) · G(Y_{i,t}) · Δ_X, where G(τ) = 2Φ(√τ) − 1 maps posterior "
    "precision to task performance (probability of correctly identifying the state). "
    "If either precision is zero, output is zero. Shared and task-specific knowledge "
    "are not substitutes — they are complements. This is the key assumption that makes "
    "AI substitution destructive.",
    space_after=6
)

heading("Agents and effort", level=3)
body(
    "A continuum of short-lived agents (one generation per period) on an island of "
    "size N choose effort e_{i,t} ≥ 0 at convex cost e^α/α. Effort simultaneously "
    "produces two signals:",
    space_after=2
)
bullet("Private signal about θ_{i,t} with precision λ_I · e (kept by the agent)")
bullet("Public signal about θ_t with precision λ_G · e per agent, aggregated to λ_G · N · ē across the island")
body(
    "The aggregate public signal precision determines X_{t+1} (next period's prior). "
    "Since each agent's share of the public good is 1/N → 0 in the continuum limit, "
    "agents do not internalise the social value of their effort. This is the fundamental externality.",
    space_after=6
)

heading("AI and the first-order condition", level=3)
body(
    "Agentic AI provides each agent a free private signal about θ_{i,t} with precision τ_A. "
    "Total task precision becomes Y_{i,t} = σ⁻² + λ_I · e + τ_A. "
    "The equilibrium effort satisfies the FOC (agents ignore public-good term):",
    space_after=2
)
p = doc.add_paragraph()
p.paragraph_format.left_indent = Inches(0.5)
p.paragraph_format.space_after = Pt(6)
r = p.add_run("G(X_t) · λ_I · g(σ⁻² + λ_I·e + τ_A)  =  e^(α−1)")
set_run(r, bold=True, size=11, color=BLUE)
body(
    "where g(τ) = G′(τ) = φ(√τ)/√τ is strictly decreasing. As τ_A rises, "
    "g shifts down → equilibrium effort e* falls → public signal shrinks → X_{t+1} falls.",
    space_after=6
)

heading("Law of motion", level=3)
body("The entire dynamics reduce to a one-dimensional map:", space_after=2)
p = doc.add_paragraph()
p.paragraph_format.left_indent = Inches(0.5)
p.paragraph_format.space_after = Pt(6)
r = p.add_run("X_{t+1}  =  F(X_t)  =  [(X_t + λ_G·N·e*(X_t, τ_A))⁻¹  +  Σ²]⁻¹")
set_run(r, bold=True, size=11, color=BLUE)
body(
    "The sequence is: existing precision X_t + new public signal → posterior precision → "
    "drift degrades this into next period's prior X_{t+1}. "
    "Fixed points of F are equilibria; stability is determined by F′.",
    space_after=8
)

# ── 2.3 Predictions ───────────────────────────────────────────────────────────
heading("2.3  Analytical Predictions", level=2)

add_table(
    headers=["Prediction", "Condition", "Result"],
    rows=[
        ["Regime bifurcation",
         "ε = 1/(α−1) vs. 4",
         "ε > 4: collapse SS stable, multiple equilibria possible. "
         "ε ≤ 4: unique stable high-knowledge SS, collapse impossible."],
        ["Path dependence",
         "Elastic regime, τ_A < τ_A^c",
         "Three SS: X=0 (stable), X̄_m (unstable tipping point), X̄_h (stable). "
         "Outcome depends on whether X_0 > X̄_m."],
        ["Monotone erosion",
         "Both regimes",
         "X̄_h strictly decreasing in τ_A. AI always erodes shared knowledge at equilibrium."],
        ["Non-monotone welfare",
         "Both regimes",
         "W(τ_A) is hump-shaped with interior maximum τ_A* < τ_A^c."],
        ["Logarithmic scaling",
         "Elastic regime",
         "τ_A^c = O(ln N). Larger communities are more resilient but protection is sub-linear."],
    ],
    col_widths=[1.5, 1.5, 3.5]
)

# ── 2.4 ABM translation ───────────────────────────────────────────────────────
heading("2.4  Translation to Agent-Based Model", level=2)
body(
    "The analytical model is a deterministic dynamical system on X_t. "
    "We translate it into a simulation where agents are explicit, signals are drawn, "
    "and belief states are tracked individually. Key mapping decisions:",
    space_after=4
)

add_table(
    headers=["Analytical element", "ABM implementation", "Fidelity"],
    rows=[
        ["Continuum of agents", "N discrete agents; public FOC ignores 1/N share (continuum limit preserved)", "Exact"],
        ["Optimal effort e*(X,τ)", "Numerical root-finding on FOC via Brentq; lower bound 1e-15 to handle near-zero X", "Exact"],
        ["Law of motion F(X)", "Signal first, then drift: X_{t+1} = [(X_t + λ_G·N·e*)⁻¹ + Σ²]⁻¹", "Exact"],
        ["Public signal aggregation", "Collective posterior precision = X_t + λ_G·N·e*", "Exact"],
        ["Steady-state detection", "Log-spaced grid down to 1e-9 + binary search for τ_A^c", "Numerical"],
        ["Agent belief visualisation", "Each agent draws their own signal (precision λ_G·e) before aggregation — creates heterogeneity", "Cosmetic extension"],
        ["σ⁻² parameter", "Set to 0.01 (not paper's 1.0) for numerical visibility; all qualitative predictions preserved", "Quantitative shift"],
    ],
    col_widths=[2.2, 3.5, 1.3]
)

# ── 2.5 Fixes applied ─────────────────────────────────────────────────────────
heading("2.5  Critical Fixes Applied During Development", level=2)
body(
    "Several non-trivial errors in the initial implementation were identified and corrected:",
    space_after=4
)

bullet(
    "Initial code applied drift before aggregating the public signal, reversing the correct "
    "order in the law of motion. The paper aggregates first, then drifts. This shifted all "
    "steady-state values and invalidated quantitative predictions.",
    bold_prefix="Law-of-motion order (critical)"
)
bullet(
    "σ⁻² = 0 broke the e*(X) ~ X^{1/(2(α−1))} scaling near X=0 that drives the ε=4 "
    "regime bifurcation. Set to σ⁻² = 0.01 to restore correct asymptotic behaviour while "
    "keeping dynamics visible.",
    bold_prefix="σ⁻² = 0 destroyed regime bifurcation (critical)"
)
bullet(
    "Original solver used lower bound 1e-12, missing very small but non-zero e* values "
    "near the collapse basin. Reduced to 1e-15.",
    bold_prefix="Solver lower bound (medium)"
)
bullet(
    "Linear scan for τ_A^c with step size 0.38 produced step-function output. Replaced "
    "with binary search to tolerance 1e-3.",
    bold_prefix="Collapse threshold scan resolution (medium)"
)
bullet(
    "Steady-state grid starting at 1e-6 missed X̄_m < 1e-6 in the elastic regime at low τ_A. "
    "Extended log grid to 1e-9.",
    bold_prefix="Steady-state grid near zero (medium)"
)
doc.add_paragraph()

# ── 2.6 Verification ─────────────────────────────────────────────────────────
heading("2.6  Prediction Verification (Default Parameters)", level=2)
body("Default: α=1.20 (ε=5.0), N=50, λ_G=1.0, λ_I=1.0, σ⁻²=0.01, Σ²=0.05", space_after=4)

add_table(
    headers=["Prediction", "Result", "Status"],
    rows=[
        ["Regime flip at ε=4", "F(1e-8)−X < 0 for α<1.25, > 0 for α≥1.25", "✓ Exact"],
        ["3 steady states at τ_A=0.30", "X = 0,  X̄_m ≈ 0.00127,  X̄_h ≈ 5.895", "✓ Reproduced"],
        ["X̄_h monotone in τ_A", "10.33 → 8.20 → 5.58 → 1.86 as τ_A: 0→0.57", "✓ Strict"],
        ["Interior welfare maximum τ_A*", "τ_A* ≈ 0.48,  τ_A^c ≈ 0.619  (interior)", "✓ Reproduced"],
        ["Scaling law τ_A^c ~ ln(N)", "R² = 0.993,  slope = 0.162·ln(N) + 0.002", "✓ Strong fit"],
    ],
    col_widths=[2.5, 3.0, 1.5]
)

# ── 2.7 UI ────────────────────────────────────────────────────────────────────
heading("2.7  Interactive Application", level=2)
body(
    "Built in Python with Streamlit. Runs at localhost:8502. "
    "All computations are cached; the UI updates live as sliders move.",
    space_after=4
)
add_table(
    headers=["Tab", "Content", "Key feature"],
    rows=[
        ["📐 Knowledge Dynamics", "F(X) cobweb plot with equilibria marked", "Live stability classification (slope test)"],
        ["📈 Time Paths X_t", "Deterministic trajectories from multiple X_0", "Low/high start shows path dependence"],
        ["💡 Social Welfare W", "W(τ_A) hump + knowledge erosion decomposition", "τ_A* and τ_A^c marked; welfare components split"],
        ["📏 Community Size Effect", "τ_A^c vs N with log-linear fit and R²", "Shading shows which N values are safe at current τ_A"],
        ["🧠 Agent Beliefs", "2D belief-error cloud, two precision ellipses", "Individual vs collective σ shows externality gap live"],
        ["📖 Model Guide", "7 expandable sections + quick-reference table", "Every symbol: plain English + formal definition + range"],
    ],
    col_widths=[1.8, 2.8, 2.4]
)

# ── 2.8 Key insight ────────────────────────────────────────────────────────────
heading("2.8  Key Emergent Insight from the Simulation", level=2)
body(
    "The Agent Beliefs tab revealed a structural feature not emphasised in the paper: "
    "the externality gap (σ_individual − σ_collective) is largest in the high-knowledge "
    "equilibrium, not during collapse. When X_t is high, agents exert substantial effort, "
    "produce a rich aggregate signal, and the collective knowledge far exceeds any individual's. "
    "The public-good problem is most severe — and most invisible — precisely when society "
    "appears to be doing well. The danger is hidden inside success.",
    space_after=8, italic=True
)

rule()

# ── 2.9 Parameter reference ───────────────────────────────────────────────────
heading("2.9  Full Parameter and Variable Reference", level=2)

add_table(
    headers=["Symbol", "Name", "Type", "Default", "Key role"],
    rows=[
        ["α",       "Effort cost steepness",       "Parameter",  "1.20",  "Regime switch: ε = 1/(α−1); collapse if ε > 4 (α < 1.25)"],
        ["ε",       "Effort elasticity",            "Derived",    "5.0",   "Critical threshold at 4"],
        ["N",       "Community size",               "Parameter",  "50",    "τ_A^c ~ ln(N)"],
        ["λ_G",     "Public learning efficiency",   "Parameter",  "1.0",   "Strength of knowledge externality"],
        ["λ_I",     "Private learning efficiency",  "Parameter",  "1.0",   "Private return to effort"],
        ["Σ²",      "Knowledge decay rate",         "Parameter",  "0.05",  "How fast shared knowledge drifts"],
        ["σ⁻²",     "Task prior precision",         "Parameter",  "0.01",  "Baseline task knowledge without effort"],
        ["τ_A",     "AI capability",                "Policy var", "0.0",   "Substitutes for private learning; key driver"],
        ["θ_t",     "Shared knowledge state",       "Exogenous",  "—",     "Random walk; never observed directly"],
        ["θ_{i,t}", "Task-specific state",          "Exogenous",  "—",     "Fresh each period; idiosyncratic"],
        ["X_t",     "Shared knowledge precision",   "State var",  "—",     "Core state variable; drives all dynamics"],
        ["e*",      "Equilibrium effort",           "Endogenous", "—",     "Set by FOC; declines in τ_A"],
        ["G(τ)",    "Knowledge quality",            "Function",   "—",     "2Φ(√τ)−1; bounded [0,1]; concave"],
        ["F(X)",    "Law of motion",                "Function",   "—",     "X_{t+1}; fixed points are equilibria"],
        ["X̄_h",    "High-knowledge equilibrium",   "Equilibrium","—",     "Stable; decreasing in τ_A"],
        ["X̄_m",    "Tipping point",                "Equilibrium","—",     "Unstable; basin boundary"],
        ["τ_A^c",   "Collapse threshold",           "Threshold",  "0.619", "Above this: no positive equilibrium"],
        ["τ_A*",    "Welfare-maximising AI level",  "Optimum",    "0.48",  "Policy target: cap τ_A here"],
        ["W",       "Social welfare",               "Welfare",    "—",     "G(X̄_h)·G(Ȳ) − ē^α/α; hump-shaped in τ_A"],
    ],
    col_widths=[0.6, 1.8, 0.9, 0.6, 3.1]
)

# ── Footer ────────────────────────────────────────────────────────────────────
doc.add_paragraph()
rule()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Knowledge Collapse Project  ·  Simulation and UI: knowledge_collapse_abm.py + knowledge_collapse_ui.py")
set_run(run, size=9, color=GREY, italic=True)

# ── Save ──────────────────────────────────────────────────────────────────────
doc.save(OUT)
print(f"Saved: {OUT}")
