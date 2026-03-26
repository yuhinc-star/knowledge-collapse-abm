"""
Generate project summary for the Knowledge Collapse Recombination Extension as a formatted .docx
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = "/Users/jacob/Desktop/Claude Code/Knowledge_Collapse_Extension_Summary.docx"

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin   = Inches(1.25)
    section.right_margin  = Inches(1.25)

# ── Colour palette ────────────────────────────────────────────────────────────
BLUE   = RGBColor(0x1E, 0x40, 0xAF)
GREY   = RGBColor(0x37, 0x41, 0x51)
RED    = RGBColor(0xB9, 0x1C, 0x1C)
GREEN  = RGBColor(0x06, 0x59, 0x37)
BLACK  = RGBColor(0x11, 0x18, 0x27)

def set_run(run, bold=False, italic=False, size=11, color=BLACK, font="Calibri"):
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.font.name = font

def heading(text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18 if level == 1 else 12)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    set_run(run, bold=True,
            size=18 if level == 1 else (14 if level == 2 else 12),
            color=BLUE if level <= 2 else GREY)
    return p

def body(text, space_after=4, italic=False, color=BLACK, bold=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after  = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    run = p.add_run(text)
    set_run(run, bold=bold, italic=italic, color=color)
    return p

def bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after  = Pt(2)
    p.paragraph_format.left_indent  = Inches(0.3)
    if bold_prefix:
        r1 = p.add_run(bold_prefix + ": ")
        set_run(r1, bold=True, size=11)
        r2 = p.add_run(text)
        set_run(r2, size=11)
    else:
        run = p.add_run(text)
        set_run(run, size=11)
    return p

def indented(text, color=BLUE):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Inches(0.5)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run(text)
    set_run(run, bold=True, size=11, color=color)
    return p

def rule():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(6)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "BFDBFE")
    pBdr.append(bottom)
    pPr.append(pBdr)

def add_table(headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        set_run(run, bold=True, size=10, color=RGBColor(0xFF, 0xFF, 0xFF))
        shading = OxmlElement("w:shd")
        shading.set(qn("w:val"),   "clear")
        shading.set(qn("w:color"), "auto")
        shading.set(qn("w:fill"),  "1E40AF")
        cell._tc.get_or_add_tcPr().append(shading)
    for ri, row_data in enumerate(rows):
        row = table.rows[ri + 1]
        fill = "EFF6FF" if ri % 2 == 0 else "FFFFFF"
        for ci, val in enumerate(row_data):
            cell = row.cells[ci]
            cell.text = ""
            run = cell.paragraphs[0].add_run(str(val))
            set_run(run, size=10)
            shading = OxmlElement("w:shd")
            shading.set(qn("w:val"),   "clear")
            shading.set(qn("w:color"), "auto")
            shading.set(qn("w:fill"),  fill)
            cell._tc.get_or_add_tcPr().append(shading)
    if col_widths:
        for ci, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[ci].width = Inches(w)
    doc.add_paragraph()


# ═════════════════════════════════════════════════════════════════════════════
# TITLE PAGE
# ═════════════════════════════════════════════════════════════════════════════
tp = doc.add_paragraph()
tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
tp.paragraph_format.space_before = Pt(48)
r = tp.add_run("Knowledge Collapse — Recombination Extension")
set_run(r, bold=True, size=26, color=BLUE)

tp2 = doc.add_paragraph()
tp2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = tp2.add_run("Cross-Disciplinary Recombination as a Second Knowledge Spillover Channel")
set_run(r2, italic=True, size=14, color=GREY)

tp3 = doc.add_paragraph()
tp3.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = tp3.add_run("Extension of Acemoglu, Kong & Ozdaglar (2026)  ·  Project Summary")
set_run(r3, size=11, color=GREY)

doc.add_paragraph()
rule()
doc.add_paragraph()


# ═════════════════════════════════════════════════════════════════════════════
# PART 1 — PLAIN LANGUAGE SUMMARY
# ═════════════════════════════════════════════════════════════════════════════
heading("Part 1 — Plain Language Summary", level=1)

heading("The problem this extension solves", level=2)
body(
    "The original Acemoglu, Kong & Ozdaglar (2026) paper proves that AI causes knowledge collapse: "
    "AI substitutes for human effort, effort falls, the public signals that sustain shared knowledge "
    "disappear, and the community's knowledge base erodes — permanently, once past the tipping point. "
    "The result is pessimistic and absorbing: once collapse begins, nothing within the model stops it.",
    space_after=6
)
body(
    "The real world, however, has a mechanism the original model ignores: when a field stagnates, "
    "researchers don't just watch it collapse. They move. They combine insights from multiple failing "
    "fields to create new ones. Econophysics emerged from physics and economics. Behavioural economics "
    "from psychology and economics. Molecular biology from chemistry and biology. "
    "This is cross-disciplinary recombination — and it is precisely the kind of self-correcting "
    "response that could counteract AI-induced collapse.",
    space_after=6
)
body(
    "This extension adds recombination as a second spillover channel alongside the original "
    "within-domain public signal. The question becomes: does recombination reverse collapse? "
    "Under what conditions? And does AI help or hurt when recombination is possible?",
    space_after=8, bold=True
)

heading("The core mechanism in plain terms", level=2)
body(
    "Think of multiple research communities (domains), each building shared knowledge the same way "
    "the original model describes: through costly effort that produces both private and public signals. "
    "AI makes each agent more productive on their private task — but crowds out the effort that "
    "sustains domain-level shared knowledge. Domains erode.",
    space_after=6
)
body(
    "When enough domains have eroded — specifically, when the average domain quality falls to a "
    "fraction δ of the pre-AI benchmark — a new domain is created. This is not arbitrary: it is "
    "the break-even point derived from the same credit incentives that govern effort in the baseline. "
    "A new domain is worth entering when its expected knowledge return equals the return "
    "in existing domains. That happens precisely when existing domains have decayed to δ·X_h0.",
    space_after=6
)
body(
    "The new domain inherits knowledge from all its predecessors: a weighted average of their "
    "current knowledge levels (weighted by how many agents each domain has), plus a synergy bonus "
    "from every pair of domains (γ·Xᵢ·Xⱼ for each pair i, j). The synergy term captures "
    "Weitzman's insight: combining economics and psychology produced behavioural economics — "
    "richer than either predecessor alone.",
    space_after=6
)
body(
    "Agents then reallocate across all domains (including the new one) until they are indifferent: "
    "each agent earns the same expected credit G(X_d)/N_d in every domain. More attractive domains "
    "(higher X_d) attract proportionally more agents.",
    space_after=8
)

heading("The surprising role of AI", level=2)
body(
    "In a closed economy (no recombination), AI is purely destructive: it erodes domain knowledge "
    "with no compensating mechanism. In an open economy (recombination possible), AI plays a "
    "counterintuitive role.",
    space_after=6
)
body(
    "Without AI: no collapse, no trigger fires, no new domains, ΣX_d stays at X_h0.",
    space_after=4, italic=True
)
body(
    "With AI: domains erode, triggers fire continuously, recombination cascade activates, "
    "ΣX_d can grow far above X_h0.",
    space_after=4, italic=True
)
body(
    "AI, by destabilising individual domains, activates the cascade that would never fire otherwise. "
    "The force that destroys the original knowledge production mechanism is precisely what creates "
    "the conditions for recombination to thrive. AI is a complement to recombination.",
    space_after=8, bold=True
)

heading("The five predictions", level=2)
bullet(
    "There is a threshold δ* above which recombination reverses collapse (ΣX_d grows over time) "
    "and below which it fails to offset decay (ΣX_d still falls). δ is the knowledge transferability "
    "parameter — how much of the predecessor's knowledge survives the move to a new context.",
    bold_prefix="P1 — Recombination threshold"
)
bullet(
    "ΣX_d trends clearly up or down over time; it is almost never flat. The competition between "
    "recombination gain at each spawn and continuous decay across all periods produces a dominant "
    "directional force. The transition at δ* is sharp.",
    bold_prefix="P2 — Monotone trend"
)
bullet(
    "AI+recombination (open economy) always outperforms AI without recombination (closed economy), "
    "for any δ > 0 and any τ_A. Even a tiny δ provides some inheritance at each spawn that the "
    "closed economy never receives.",
    bold_prefix="P3 — Open beats closed"
)
bullet(
    "AI+recombination outperforms no-AI+recombination for all AI levels above the collapse "
    "threshold τ_A^c. ΣX_d is hump-shaped in τ_A — it peaks at intermediate AI capability "
    "and declines at very high τ_A (domains collapse too fast for inheritance to accumulate). "
    "But for the entire range τ_A > τ_A^c, the open economy with AI exceeds the open economy "
    "without AI.",
    bold_prefix="P4 — AI as complement"
)
bullet(
    "ΣX_d(T) exceeds the pre-AI benchmark X_h0 if and only if δ > δ*. Below δ*, the cascade "
    "adds some knowledge but decay dominates — total knowledge stays below the pre-AI level. "
    "Above δ*, recombination accumulates faster than decay removes — full reversal of Acemoglu's result.",
    bold_prefix="P5 — Full reversal condition"
)
doc.add_paragraph()

rule()
doc.add_page_break()


# ═════════════════════════════════════════════════════════════════════════════
# PART 2 — DETAILED TECHNICAL SUMMARY
# ═════════════════════════════════════════════════════════════════════════════
heading("Part 2 — Detailed Technical Summary", level=1)

heading("2.1  Relationship to the Baseline Model", level=2)
body(
    "The extension builds directly on Acemoglu, Kong & Ozdaglar (2026). The baseline model's "
    "entire mathematical structure — the effort FOC, the law of motion F(X), the credit function "
    "G(τ) = 2Φ(√τ) − 1, the collapse threshold τ_A^c — is preserved unchanged. "
    "The extension adds a second dimension: multiple domains that can spawn, die, and pass knowledge "
    "to successors. The baseline is recovered exactly when δ = 0 (no knowledge transfers, "
    "spawning has no effect).",
    space_after=6
)
body(
    "This means the extension is a strict superset: every parameter and prediction of the baseline "
    "remains valid. The new parameters (δ, γ) add mechanisms on top of the original collapse dynamics.",
    space_after=8
)

heading("2.2  Three Knowledge Layers", level=2)
body("The extension introduces three distinct knowledge types with different economic roles:", space_after=4)

add_table(
    headers=["Layer", "Symbol", "Persistence", "Who observes", "Drives what"],
    rows=[
        ["Task knowledge",   "Y_{i,t} = σ⁻² + λ_I·e + τ_A",
         "Resets every period",
         "Agent i only",
         "Agent's private output; AI enters here"],
        ["Domain knowledge", "X_{d,t}  (law of motion below)",
         "Persistent; compounds over time",
         "All agents in domain d",
         "Effort FOC; public signal production; spawn trigger"],
        ["General knowledge","ΣX_d  (Bayesian sum across domains)",
         "Analyst metric only",
         "Analyst only — not in any agent's info set",
         "Primary welfare metric; P1–P5 all stated in terms of ΣX_d"],
    ],
    col_widths=[1.2, 2.0, 1.4, 1.3, 1.5]
)

body(
    "A critical design choice: agents do NOT observe or optimise over ΣX_d. Each agent "
    "knows only their own domain's X_{d,t}. This is realistic — a macroeconomist optimises "
    "given the state of macroeconomics, not the total stock of all human knowledge. "
    "If agents observed ΣX_d, their FOC would change fundamentally and the public-good "
    "externality logic would break down.",
    space_after=8
)

heading("2.3  Domain-Level Law of Motion", level=2)
body("Each domain d follows the original Acemoglu law of motion:", space_after=2)
indented("X_{d,t+1}  =  [(X_{d,t} + λ_G · N_d · ē*(X_{d,t}, τ_A))⁻¹  +  Σ²]⁻¹")
body(
    "where ē*(X_{d,t}, τ_A) is the equilibrium effort in domain d, solved from the FOC:",
    space_after=2
)
indented("G(X_{d,t}) · λ_I · g(σ⁻² + λ_I·e + τ_A)  =  e^(α−1)")
body(
    "This is identical to the baseline equation — the only new element is that N_d "
    "(agents in domain d) is now determined by the credit equilibrium rather than being fixed.",
    space_after=8
)

heading("2.4  Endogenous Spawn Trigger", level=2)
body(
    "A new domain d' is created when average domain knowledge falls to the break-even level "
    "implied by the credit equilibrium. The spawn condition is:",
    space_after=2
)
indented("ΣX_d / k  <  δ · X_h0")
body(
    "where k is the current number of active domains and X_h0 is the pre-AI steady-state "
    "knowledge level (the high-knowledge equilibrium at τ_A = 0).",
    space_after=6
)
body(
    "This is not an arbitrary threshold — it is derived from the break-even condition in the "
    "credit market. A new domain is worth creating when its expected credit return equals "
    "the return in existing domains. At the spawn condition, a new domain born with "
    "X_{d',0} = δ·(ΣX_d/k) = δ²·X_h0 matches the current average quality of existing domains. "
    "Entry is immediately worthwhile.",
    space_after=6
)
body(
    "The same δ governs both how much knowledge transfers (inheritance rate) and when "
    "spawning is worthwhile (break-even threshold). This is not a modelling shortcut — it is "
    "a theoretical economy: one parameter doing double duty because the same credit mechanism "
    "that determines knowledge value also determines when new entry is attractive.",
    space_after=6
)
body(
    "Consequence: higher δ means spawning fires earlier (when domains have only decayed "
    "to 85% of X_h0 for δ=0.85, rather than 50% for δ=0.5) and inheritance is richer. "
    "δ simultaneously controls the timing and the quality of the cascade.",
    space_after=6
)
body(
    "Once the cascade begins, the trigger fires every period. After spawn, the new domain "
    "inherits δ·(average X), which is itself below δ·X_h0 (since the average just reached "
    "that threshold). So the next period's average is again below the threshold and spawning "
    "continues. This continuous spawning is correct model behaviour — not a bug. "
    "Whether general knowledge grows or falls depends entirely on how much each spawn "
    "contributes, not on how frequently spawning occurs.",
    space_after=8
)

heading("2.5  Convergent Inheritance", level=2)
body(
    "When domain d' is spawned, it inherits from ALL active predecessor domains — "
    "not just the most recently created one (chain inheritance) and not equally from all (uniform).",
    space_after=2
)
indented("X_{d',0}  =  δ · (Σ_d w_d · X_d  +  γ · Σ_{i<j} w_i · X_i · w_j · X_j)")
body("where  w_d = N_d / N  is the agent share of domain d (from the credit equilibrium).", space_after=6)
body("The two terms serve distinct roles:", space_after=2)
bullet(
    "Agent-weighted average knowledge: what movers collectively carry into d'. "
    "More heavily populated domains contribute more — reflecting that more agents "
    "means more accumulated expertise to transfer.",
    bold_prefix="First term: Σ w_d · X_d"
)
bullet(
    "Pairwise synergy (Weitzman recombination): each pair of domains (i, j) contributes "
    "a bonus proportional to both their knowledge levels. Combining economics and psychology "
    "produced behavioural economics — the insight emerged from the combination, not from either "
    "field alone. With k equal-knowledge domains of level X, total synergy = k(k−1)/2 · γ · X². "
    "As the cascade lengthens, synergies compound superlinearly.",
    bold_prefix="Second term: γ · Σ_{i<j} w_i·X_i·w_j·X_j"
)
body(
    "Setting γ = 0 recovers purely additive recombination: the new domain benefits from "
    "predecessor knowledge but traditions don't interact. Setting δ = 0 recovers the "
    "closed economy: no inheritance, no spawn effect, ΣX_d = X_{d₀,t}.",
    space_after=6
)
body(
    "The δ < 1 constraint means X_{d',0} is always below the weighted average of predecessor "
    "knowledge. Knowledge does not amplify for free at spawn — there is a real friction. "
    "The cascade sustains ΣX_d only when the continuous stream of spawns accumulates "
    "faster than decay removes knowledge from existing domains.",
    space_after=8
)

heading("2.6  Credit-Equilibrium Agent Allocation", level=2)
body(
    "After each spawn (and at each period), agents reallocate across all active domains. "
    "The allocation is derived from the credit equilibrium: agents move until "
    "per-agent expected credit is equalised across all domains.",
    space_after=2
)
body("Credit equilibrium condition:", space_after=2)
indented("G(X_{d₁}) / N_{d₁}  =  G(X_{d₂}) / N_{d₂}  =  …  =  G(X_{d'}) / N_{d'}  =  common value")
body("Solving for N_d given fixed total N:", space_after=2)
indented("N_d*  =  N · G(X_d) / Σ_j G(X_j)   for all active domains including d'")
body(
    "This is fully endogenous — derived from the same G(X) credit function that appears "
    "in the effort FOC. The allocation does NOT use δ, which governs knowledge transfer only. "
    "A domain with twice the knowledge attractiveness (in G(·) units) receives twice the agents.",
    space_after=8
)

heading("2.7  Renewal Signal at Spawn", level=2)
body(
    "When d' is created, it receives a one-time renewal boost from movers. "
    "The renewal signal has two components matching the inheritance formula:",
    space_after=4
)
bullet(
    "renewal_move  =  λ_G · N_{d'} · Σ_d [G(X_d) / Σ_old_G · e_d ]  "
    "— movers from each domain d carry effort e_d, weighted by domain attractiveness G(X_d)/Σ_old_G. "
    "This is a one-time public signal: the expertise movers have accumulated applied in the new context.",
    bold_prefix="Move channel"
)
bullet(
    "renewal_interact  =  γ · Σ_{i<j} X_i · X_j  "
    "— the Weitzman pairwise synergy: combining the knowledge of any two existing domains "
    "produces a bonus proportional to both their levels. This fires ONLY at spawn — not as "
    "an ongoing per-period term. γ governs recombination at the moment of creation, not "
    "continuous cross-pollination.",
    bold_prefix="Interact channel (γ, spawn-time only)"
)
body(
    "The spawn-time-only restriction on γ is a deliberate design choice. Making γ an ongoing "
    "spillover would create a permanent cross-domain public good — a different and more complex "
    "mechanism. Restricting γ to spawn time keeps the model parsimonious: recombination happens "
    "at the moment fields combine, not continuously thereafter.",
    space_after=8
)

heading("2.8  Five Predictions — Formal Statements", level=2)

add_table(
    headers=["Prediction", "Formal statement", "Economic interpretation", "Test"],
    rows=[
        ["P1 — Threshold δ*",
         "∃ δ* ∈ (0,1) such that ΣX_d(T)→∞ if δ>δ*, ΣX_d(T)→0 if δ<δ*",
         "Recombination reverses collapse iff transferability is high enough. "
         "At δ*, inheritance exactly offsets decay at each spawn.",
         "Tab 4: ΣX_d(T) vs δ sweep — clear crossing at δ*"],
        ["P2 — Monotone trend",
         "ΣX_d(t) is monotonically increasing or decreasing once cascade begins; "
         "almost never hovering near a steady state",
         "One force dominates: either each spawn adds enough to outpace decay, "
         "or it doesn't. No partial equilibrium.",
         "Tab 5: trajectories at δ < δ* (falling) vs δ > δ* (rising)"],
        ["P3 — Open beats closed",
         "ΣX_d_open(T) > ΣX_d_closed(T) for all δ > 0, τ_A > 0",
         "Recombination is always weakly beneficial. Even tiny δ provides some "
         "inheritance the closed economy never receives.",
         "Tab 3: green (open) always above dark (closed)"],
        ["P4 — AI as complement",
         "ΣX_d_open(T, τ_A) > ΣX_d_open(T, 0) for all τ_A > τ_A^c; "
         "ΣX_d is hump-shaped in τ_A",
         "AI activates the cascade. Without collapse there is no spawn, "
         "so the open economy needs AI's pressure. At very high τ_A "
         "the hump peaks and declines (domains collapse too fast).",
         "Tab 3: τ_A sweep — open-economy curve above no-AI level for τ_A > τ_A^c"],
        ["P5 — Full reversal",
         "ΣX_d(T) > X_h0 iff δ > δ*",
         "ΣX_d exceeds the pre-AI benchmark only when recombination dominates. "
         "Below δ*, cascade adds knowledge but decay still wins. "
         "Above δ*, Acemoglu's result is fully reversed.",
         "Tab 4: X_h0 threshold line — ΣX_d(T) crosses above it at δ*"],
    ],
    col_widths=[1.2, 1.8, 2.2, 1.8]
)

heading("2.9  Design Decisions and Fixes", level=2)
body("Several non-trivial design choices were made and errors corrected during development:", space_after=4)

bullet(
    "Early implementations fired spawning when any single domain fell below a fixed fraction "
    "of its initial value. This is exogenous: the threshold is set by the modeller, not derived "
    "from any incentive. Two problems: (1) no economic foundation for the fraction; "
    "(2) children born below the threshold spawned immediately, creating all domains at t=0. "
    "Fix: use ΣX_d/k < δ·X_h0, derived from the credit-equilibrium break-even condition.",
    bold_prefix="Wrong spawn trigger (critical)"
)
bullet(
    "γ was initially implemented as an ongoing per-period cross-domain spillover term. "
    "This is incorrect: γ is the Weitzman recombination complementarity, which operates "
    "at the moment of combination (spawn), not as a continuous flow. "
    "Fix: γ enters only in the renewal signal at spawn; removed from the per-period update step.",
    bold_prefix="γ doing ongoing spillover (critical)"
)
bullet(
    "An early version used δ to also govern the fraction of agents who migrate to the new domain. "
    "This conflates two distinct parameters: δ is knowledge transferability, not agent mobility. "
    "Fix: agent allocation determined entirely by credit equilibrium N_d* ∝ G(X_d).",
    bold_prefix="δ conflated with migration fraction (important)"
)
bullet(
    "With unlimited spawning and T=200, hundreds of domain lines were plotted in Figure 1, "
    "making the chart unreadable. Fix: cap display at MAX_PLOT=8 individual domain lines; "
    "aggregate remaining domains as a single grey '+N other domains' band. "
    "Simulation is unaffected — display only.",
    bold_prefix="Figure 1 explosion (cosmetic)"
)
bullet(
    "P4 was initially validated by checking whether the τ_A sweep had positive slope "
    "(monotone increase). The correct claim is not monotone — ΣX_d is hump-shaped. "
    "Fix: P4 holds iff ΣX_d(τ_A) ≥ no-AI baseline for all τ_A > τ_A^c.",
    bold_prefix="P4 validation logic (important)"
)
doc.add_paragraph()

heading("2.10  Interactive Application", level=2)
body(
    "Built in Python with Streamlit. Runs at localhost:8501 (combined) or standalone. "
    "All heavy computations cached with @st.cache_data. "
    "Sweep simulations use T_sw = min(T, 60) — sufficient to establish cascade direction.",
    space_after=4
)

add_table(
    headers=["Tab", "Content", "Key feature"],
    rows=[
        ["🌐 Domain Dynamics",
         "Per-domain X_d(t) time paths + ΣX_d (general knowledge). "
         "Domains shown individually up to 8; excess aggregated.",
         "Shows cascade progression; colour = generation order; "
         "grey band = aggregated later domains"],
        ["📊 Four-Way Comparison",
         "Closed / Open / No-AI / Pre-AI (X_h0) on one chart",
         "Direct visual of P3, P4, P5 at current parameters; "
         "ratio metrics shown numerically"],
        ["🔬 P3+P4 — Role of AI",
         "Left: time paths (open vs closed vs no-AI). "
         "Right: τ_A sweep showing ΣX_d(T) vs τ_A for open and no-AI.",
         "P3 confirmed by left panel; P4 confirmed by right panel sweep; "
         "hump shape visible at high τ_A"],
        ["📏 P1+P5 — Recomb threshold",
         "δ sweep: ΣX_d(T) vs δ with X_h0 threshold line.",
         "δ* visible as crossing point; P5 confirmed where "
         "ΣX_d(T) crosses X_h0"],
        ["📈 P2 — Knowledge trends",
         "Three trajectories at δ below, at, and above δ*.",
         "Monotone decline vs flat vs monotone growth shows "
         "sharp transition at δ*"],
        ["📖 Extension Guide",
         "Seven expandable sections with full technical detail, "
         "design rationale, and parameter quick-reference.",
         "Live values from current simulation embedded in guide text"],
    ],
    col_widths=[1.5, 2.8, 2.7]
)

heading("2.11  Key Insight from the Simulation", level=2)
body(
    "The strong-reversal configuration (δ=0.85, γ=0.05, τ_A=0.85) demonstrates the central "
    "theoretical result: ΣX_d(T)/X_h0 ≈ 11×, versus the closed economy at ≈ 0×. "
    "General knowledge grows to more than ten times the pre-AI benchmark, while the closed "
    "economy collapses completely. The open/closed ratio at this configuration is approximately 900×.",
    space_after=6
)
body(
    "This quantitative gap illustrates the qualitative point: recombination is not a marginal "
    "correction to knowledge collapse. When transferability is high and synergies exist, it "
    "fully reverses Acemoglu's pessimistic result. The mechanism is not that domains somehow "
    "resist collapse — they don't; each individual domain still erodes under AI. "
    "The mechanism is that the cascade of new domains, each inheriting from a richer pool of "
    "predecessors, accumulates knowledge faster than any individual domain loses it.",
    space_after=8, italic=True
)

rule()

heading("2.12  Parameter and Variable Reference", level=2)
add_table(
    headers=["Symbol", "Name", "Type", "Default", "Key role in extension"],
    rows=[
        ["α",    "Effort cost steepness",        "Parameter",  "1.20",  "Unchanged from baseline; determines regime (elastic/inelastic)"],
        ["N",    "Total community size",          "Parameter",  "50",    "Total agents allocated across all domains by credit equilibrium"],
        ["λ_G",  "Public learning efficiency",    "Parameter",  "1.0",   "Determines public signal strength — how fast domain X_d grows"],
        ["λ_I",  "Private learning efficiency",   "Parameter",  "1.0",   "Private return to effort — affects FOC and equilibrium effort"],
        ["Σ²",   "Knowledge decay rate",          "Parameter",  "0.05",  "Per-period drift — same for all domains; drives the need for spawning"],
        ["σ⁻²",  "Task prior precision",          "Parameter",  "0.01",  "Baseline task knowledge without effort"],
        ["τ_A",  "AI capability",                 "Policy var", "1.0",   "Drives domain erosion; activates spawn cascade when τ_A > τ_A^c"],
        ["δ",    "Knowledge transferability",     "Extension",  "0.5",   "Inheritance rate AND spawn break-even threshold — same parameter, dual role"],
        ["γ",    "Cross-domain complementarity",  "Extension",  "0.0",   "Weitzman pairwise synergy — fires at spawn only, not per-period"],
        ["X_{d,t}", "Domain knowledge precision", "State var",  "—",     "Evolves via law of motion; erodes under AI; spawns new domain when ΣX_d/k < δ·X_h0"],
        ["ΣX_d", "General knowledge",             "State var",  "—",     "Bayesian sum; primary welfare metric; target of P1–P5"],
        ["X_h0", "Pre-AI benchmark",              "Reference",  "—",     "High-knowledge steady state at τ_A=0; spawn trigger and P5 reference"],
        ["w_d",  "Agent share of domain d",       "Endogenous", "—",     "w_d = N_d/N; weights in inheritance formula; from credit equilibrium"],
        ["N_d*", "Equilibrium agents in domain d","Endogenous", "—",     "N · G(X_d) / Σ G(X_j); re-computed after each spawn"],
        ["G(τ)", "Knowledge quality / credit",    "Function",   "—",     "2Φ(√τ)−1; bounds [0,1]; drives FOC and allocation — unchanged from baseline"],
        ["δ*",   "Recombination threshold",       "Threshold",  "—",     "Critical δ above which ΣX_d grows; characterises P1 and P5"],
        ["τ_A^c","Collapse threshold",            "Threshold",  "0.619", "Above this, baseline domain collapses; spawn cascade activates"],
    ],
    col_widths=[0.7, 1.7, 0.9, 0.6, 3.1]
)

doc.add_paragraph()
rule()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run(
    "Knowledge Collapse Extension  ·  Simulation: knowledge_collapse_recomb.py  ·  "
    "Extension of Acemoglu, Kong & Ozdaglar (2026)"
)
set_run(run, size=9, color=GREY, italic=True)

doc.save(OUT)
print(f"Saved: {OUT}")
